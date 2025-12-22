"""Document loaders for various file formats.

This module provides abstract and concrete document loader implementations
for PDF, DOCX, HTML, and Markdown files.
"""

from abc import ABC, abstractmethod
from pathlib import Path
from typing import List, Optional, Dict, Any
import re

from .models import Document, DocumentMetadata


class DocumentLoader(ABC):
    """Abstract base class for document loaders.
    
    Document loaders are responsible for reading files of specific formats
    and converting them into Document objects with extracted text content
    and metadata.
    
    Subclasses must implement the load() method to handle their specific
    file format.
    
    Attributes:
        supported_extensions: List of file extensions this loader supports
    """
    
    supported_extensions: List[str] = []
    
    @abstractmethod
    def load(self, path: Path) -> List[Document]:
        """Load a document from the given file path.
        
        Args:
            path: Path to the file to load
            
        Returns:
            List of Document objects extracted from the file.
            May return multiple documents (e.g., one per page for PDFs).
            
        Raises:
            FileNotFoundError: If the file does not exist
            ValueError: If the file format is not supported
            IOError: If the file cannot be read
        """
        ...
    
    def can_load(self, path: Path) -> bool:
        """Check if this loader can handle the given file.
        
        Args:
            path: Path to check
            
        Returns:
            True if this loader supports the file's extension
        """
        return path.suffix.lower() in self.supported_extensions
    
    def _validate_path(self, path: Path) -> None:
        """Validate that the path exists and is supported.
        
        Args:
            path: Path to validate
            
        Raises:
            FileNotFoundError: If the file does not exist
            ValueError: If the file format is not supported
        """
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")
        if not self.can_load(path):
            raise ValueError(
                f"Unsupported file format: {path.suffix}. "
                f"Supported formats: {self.supported_extensions}"
            )



class PDFLoader(DocumentLoader):
    """Document loader for PDF files.
    
    Uses pypdf for basic text extraction and pdfplumber for enhanced
    extraction when available. Falls back gracefully if pdfplumber
    is not installed.
    
    Attributes:
        extract_per_page: If True, returns one Document per page.
                         If False, returns a single Document with all content.
        use_pdfplumber: If True, attempts to use pdfplumber for better extraction.
    """
    
    supported_extensions = [".pdf"]
    
    def __init__(
        self,
        extract_per_page: bool = True,
        use_pdfplumber: bool = True
    ) -> None:
        """Initialize the PDF loader.
        
        Args:
            extract_per_page: Whether to create separate documents per page
            use_pdfplumber: Whether to use pdfplumber for enhanced extraction
        """
        self.extract_per_page = extract_per_page
        self.use_pdfplumber = use_pdfplumber
    
    def load(self, path: Path) -> List[Document]:
        """Load a PDF document.
        
        Args:
            path: Path to the PDF file
            
        Returns:
            List of Document objects (one per page if extract_per_page=True)
            
        Raises:
            FileNotFoundError: If the file does not exist
            ValueError: If the file is not a PDF
            ImportError: If required PDF libraries are not installed
        """
        self._validate_path(path)
        
        # Try pdfplumber first for better extraction
        if self.use_pdfplumber:
            try:
                return self._load_with_pdfplumber(path)
            except ImportError:
                pass  # Fall back to pypdf
        
        return self._load_with_pypdf(path)
    
    def _load_with_pypdf(self, path: Path) -> List[Document]:
        """Load PDF using pypdf library."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError(
                "pypdf is required for PDF loading. "
                "Install it with: pip install pypdf"
            )
        
        reader = PdfReader(str(path))
        total_pages = len(reader.pages)
        documents: List[Document] = []
        
        # Extract metadata from PDF
        pdf_metadata = reader.metadata or {}
        title = pdf_metadata.get("/Title")
        author = pdf_metadata.get("/Author")
        
        if self.extract_per_page:
            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                metadata = DocumentMetadata(
                    source=str(path),
                    page_number=page_num,
                    total_pages=total_pages,
                    title=title,
                    author=author,
                    file_type="pdf",
                )
                documents.append(Document(content=text, metadata=metadata))
        else:
            all_text = "\n\n".join(
                page.extract_text() or "" for page in reader.pages
            )
            metadata = DocumentMetadata(
                source=str(path),
                total_pages=total_pages,
                title=title,
                author=author,
                file_type="pdf",
            )
            documents.append(Document(content=all_text, metadata=metadata))
        
        return documents
    
    def _load_with_pdfplumber(self, path: Path) -> List[Document]:
        """Load PDF using pdfplumber library for enhanced extraction."""
        try:
            import pdfplumber
        except ImportError:
            raise ImportError(
                "pdfplumber is required for enhanced PDF loading. "
                "Install it with: pip install pdfplumber"
            )
        
        documents: List[Document] = []
        
        with pdfplumber.open(str(path)) as pdf:
            total_pages = len(pdf.pages)
            
            # Extract metadata
            pdf_metadata = pdf.metadata or {}
            title = pdf_metadata.get("Title")
            author = pdf_metadata.get("Author")
            
            if self.extract_per_page:
                for page_num, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text() or ""
                    metadata = DocumentMetadata(
                        source=str(path),
                        page_number=page_num,
                        total_pages=total_pages,
                        title=title,
                        author=author,
                        file_type="pdf",
                        extra={"width": page.width, "height": page.height},
                    )
                    documents.append(Document(content=text, metadata=metadata))
            else:
                all_text = "\n\n".join(
                    page.extract_text() or "" for page in pdf.pages
                )
                metadata = DocumentMetadata(
                    source=str(path),
                    total_pages=total_pages,
                    title=title,
                    author=author,
                    file_type="pdf",
                )
                documents.append(Document(content=all_text, metadata=metadata))
        
        return documents



class DOCXLoader(DocumentLoader):
    """Document loader for Microsoft Word DOCX files.
    
    Uses python-docx library to extract text content from Word documents,
    including paragraphs and tables.
    
    Attributes:
        include_tables: If True, extracts text from tables as well.
        preserve_structure: If True, adds section markers for headings.
    """
    
    supported_extensions = [".docx"]
    
    def __init__(
        self,
        include_tables: bool = True,
        preserve_structure: bool = True
    ) -> None:
        """Initialize the DOCX loader.
        
        Args:
            include_tables: Whether to include table content
            preserve_structure: Whether to preserve document structure markers
        """
        self.include_tables = include_tables
        self.preserve_structure = preserve_structure
    
    def load(self, path: Path) -> List[Document]:
        """Load a DOCX document.
        
        Args:
            path: Path to the DOCX file
            
        Returns:
            List containing a single Document with the extracted content
            
        Raises:
            FileNotFoundError: If the file does not exist
            ValueError: If the file is not a DOCX
            ImportError: If python-docx is not installed
        """
        self._validate_path(path)
        
        try:
            from docx import Document as DocxDocument
            from docx.opc.exceptions import PackageNotFoundError
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX loading. "
                "Install it with: pip install python-docx"
            )
        
        try:
            doc = DocxDocument(str(path))
        except PackageNotFoundError:
            raise ValueError(f"Invalid or corrupted DOCX file: {path}")
        
        # Extract core properties (metadata)
        core_props = doc.core_properties
        title = core_props.title if core_props.title else None
        author = core_props.author if core_props.author else None
        created = core_props.created
        modified = core_props.modified
        
        # Extract text content
        content_parts: List[str] = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Add structure markers for headings
                if self.preserve_structure and para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    try:
                        level_num = int(level)
                        prefix = "#" * level_num + " "
                    except ValueError:
                        prefix = "# "
                    content_parts.append(f"{prefix}{text}")
                else:
                    content_parts.append(text)
        
        # Extract table content
        if self.include_tables:
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    content_parts.append(table_text)
        
        content = "\n\n".join(content_parts)
        
        metadata = DocumentMetadata(
            source=str(path),
            title=title,
            author=author,
            created_date=created,
            modified_date=modified,
            file_type="docx",
            extra={
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            },
        )
        
        return [Document(content=content, metadata=metadata)]
    
    def _extract_table_text(self, table: Any) -> str:
        """Extract text from a table as formatted text.
        
        Args:
            table: A python-docx Table object
            
        Returns:
            Formatted string representation of the table
        """
        rows: List[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)



class HTMLLoader(DocumentLoader):
    """Document loader for HTML files.
    
    Uses BeautifulSoup to parse HTML and extract text content,
    with options to preserve structure and filter elements.
    
    Attributes:
        remove_scripts: If True, removes script and style tags.
        preserve_links: If True, converts links to markdown format.
        extract_title: If True, extracts title from <title> tag.
    """
    
    supported_extensions = [".html", ".htm"]
    
    def __init__(
        self,
        remove_scripts: bool = True,
        preserve_links: bool = False,
        extract_title: bool = True
    ) -> None:
        """Initialize the HTML loader.
        
        Args:
            remove_scripts: Whether to remove script/style elements
            preserve_links: Whether to preserve links in markdown format
            extract_title: Whether to extract title from HTML
        """
        self.remove_scripts = remove_scripts
        self.preserve_links = preserve_links
        self.extract_title = extract_title
    
    def load(self, path: Path) -> List[Document]:
        """Load an HTML document.
        
        Args:
            path: Path to the HTML file
            
        Returns:
            List containing a single Document with the extracted content
            
        Raises:
            FileNotFoundError: If the file does not exist
            ValueError: If the file is not HTML
            ImportError: If BeautifulSoup is not installed
        """
        self._validate_path(path)
        
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError(
                "beautifulsoup4 is required for HTML loading. "
                "Install it with: pip install beautifulsoup4"
            )
        
        # Read the HTML file
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, "html.parser")
        
        # Extract title
        title = None
        if self.extract_title:
            title_tag = soup.find("title")
            if title_tag:
                title = title_tag.get_text(strip=True)
        
        # Extract meta information
        author = None
        meta_author = soup.find("meta", attrs={"name": "author"})
        if meta_author:
            author = meta_author.get("content")
        
        # Remove unwanted elements
        if self.remove_scripts:
            for element in soup(["script", "style", "noscript"]):
                element.decompose()
        
        # Extract text content
        if self.preserve_links:
            content = self._extract_with_links(soup)
        else:
            content = soup.get_text(separator="\n", strip=True)
        
        # Clean up excessive whitespace
        content = re.sub(r"\n{3,}", "\n\n", content)
        
        metadata = DocumentMetadata(
            source=str(path),
            title=title,
            author=author,
            file_type="html",
            extra={
                "encoding": "utf-8",
            },
        )
        
        return [Document(content=content, metadata=metadata)]
    
    def _extract_with_links(self, soup: Any) -> str:
        """Extract text while preserving links in markdown format.
        
        Args:
            soup: BeautifulSoup object
            
        Returns:
            Text content with links in markdown format
        """
        # Convert links to markdown format
        for a_tag in soup.find_all("a", href=True):
            link_text = a_tag.get_text(strip=True)
            href = a_tag["href"]
            if link_text and href:
                a_tag.replace_with(f"[{link_text}]({href})")
        
        return soup.get_text(separator="\n", strip=True)



class MarkdownLoader(DocumentLoader):
    """Document loader for Markdown files with frontmatter support.
    
    Parses Markdown files and extracts YAML frontmatter metadata
    if present. Supports standard Markdown syntax.
    
    Attributes:
        parse_frontmatter: If True, parses YAML frontmatter.
        remove_frontmatter: If True, removes frontmatter from content.
    """
    
    supported_extensions = [".md", ".markdown"]
    
    def __init__(
        self,
        parse_frontmatter: bool = True,
        remove_frontmatter: bool = True
    ) -> None:
        """Initialize the Markdown loader.
        
        Args:
            parse_frontmatter: Whether to parse YAML frontmatter
            remove_frontmatter: Whether to remove frontmatter from content
        """
        self.parse_frontmatter = parse_frontmatter
        self.remove_frontmatter = remove_frontmatter
    
    def load(self, path: Path) -> List[Document]:
        """Load a Markdown document.
        
        Args:
            path: Path to the Markdown file
            
        Returns:
            List containing a single Document with the content
            
        Raises:
            FileNotFoundError: If the file does not exist
            ValueError: If the file is not Markdown
        """
        self._validate_path(path)
        
        # Read the file content
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        
        # Extract frontmatter if present
        frontmatter: Dict[str, Any] = {}
        if self.parse_frontmatter:
            frontmatter, content = self._extract_frontmatter(content)
        
        # Extract title from frontmatter or first heading
        title = frontmatter.get("title")
        if not title:
            title = self._extract_first_heading(content)
        
        author = frontmatter.get("author")
        
        # Handle date fields from frontmatter
        created_date = None
        if "date" in frontmatter:
            created_date = self._parse_date(frontmatter["date"])
        
        metadata = DocumentMetadata(
            source=str(path),
            title=title,
            author=author,
            created_date=created_date,
            file_type="markdown",
            extra={"frontmatter": frontmatter} if frontmatter else {},
        )
        
        return [Document(content=content.strip(), metadata=metadata)]
    
    def _extract_frontmatter(self, content: str) -> tuple[Dict[str, Any], str]:
        """Extract YAML frontmatter from content.
        
        Args:
            content: The full file content
            
        Returns:
            Tuple of (frontmatter dict, remaining content)
        """
        # Check for frontmatter delimiter
        if not content.startswith("---"):
            return {}, content
        
        # Find the closing delimiter
        lines = content.split("\n")
        end_index = -1
        for i, line in enumerate(lines[1:], start=1):
            if line.strip() == "---":
                end_index = i
                break
        
        if end_index == -1:
            return {}, content
        
        # Extract frontmatter YAML
        frontmatter_lines = lines[1:end_index]
        frontmatter_text = "\n".join(frontmatter_lines)
        
        # Parse YAML
        frontmatter = self._parse_yaml(frontmatter_text)
        
        # Get remaining content
        if self.remove_frontmatter:
            remaining_content = "\n".join(lines[end_index + 1:])
        else:
            remaining_content = content
        
        return frontmatter, remaining_content
    
    def _parse_yaml(self, yaml_text: str) -> Dict[str, Any]:
        """Parse YAML frontmatter text.
        
        Uses PyYAML if available, otherwise falls back to simple parsing.
        
        Args:
            yaml_text: YAML formatted text
            
        Returns:
            Dictionary of parsed values
        """
        try:
            import yaml
            return yaml.safe_load(yaml_text) or {}
        except ImportError:
            # Simple fallback parser for basic key: value pairs
            result: Dict[str, Any] = {}
            for line in yaml_text.split("\n"):
                line = line.strip()
                if ":" in line and not line.startswith("#"):
                    key, value = line.split(":", 1)
                    key = key.strip()
                    value = value.strip()
                    # Remove quotes if present
                    if value.startswith('"') and value.endswith('"'):
                        value = value[1:-1]
                    elif value.startswith("'") and value.endswith("'"):
                        value = value[1:-1]
                    result[key] = value
            return result
    
    def _extract_first_heading(self, content: str) -> Optional[str]:
        """Extract the first heading from Markdown content.
        
        Args:
            content: Markdown content
            
        Returns:
            The first heading text, or None if not found
        """
        # Match ATX-style headings (# Heading)
        match = re.search(r"^#+\s+(.+)$", content, re.MULTILINE)
        if match:
            return match.group(1).strip()
        return None
    
    def _parse_date(self, date_value: Any) -> Optional[Any]:
        """Parse a date value from frontmatter.
        
        Args:
            date_value: Date value (string or datetime)
            
        Returns:
            Parsed datetime or None
        """
        from datetime import datetime
        
        if isinstance(date_value, datetime):
            return date_value
        
        if isinstance(date_value, str):
            # Try common date formats
            formats = [
                "%Y-%m-%d",
                "%Y-%m-%dT%H:%M:%S",
                "%Y-%m-%d %H:%M:%S",
                "%B %d, %Y",
            ]
            for fmt in formats:
                try:
                    return datetime.strptime(date_value, fmt)
                except ValueError:
                    continue
        
        return None
